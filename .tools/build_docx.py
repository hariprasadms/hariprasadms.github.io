#!/usr/bin/env python3
"""Build a Word (.docx) interior of the book for KDP paperback.

Run with the venv that has python-docx:
    .venv-docx/bin/python .tools/build_docx.py
Reuses the chapter parsing + front/back-matter copy from the other builders.
Output: dist/from-bugs-to-brilliance.docx  (5x8in, headings, TOC field, page numbers).
"""
import os, sys, glob, re
from html.parser import HTMLParser

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
sys.path.insert(0, HERE)
from build_epub import load_chapter, META, DISCLAIMER, BIO
from build_print import DEDICATION, ACK, GLOSSARY

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(ROOT, "dist", "from-bugs-to-brilliance.docx")
SKIP = {"nlabel", "dlabel", "badge", "suite-head", "page-marker", "cb-num"}
GREEN = RGBColor(0x1F, 0x8A, 0x5B)
GRAY = RGBColor(0x66, 0x66, 0x66)

class Walker(HTMLParser):
    """Turn a chapter body into a flat list of ('para', style, runs) / ('image', path) / ('scene',) ops."""
    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.ops = []
        self.stack = []          # (tag, classes)
        self.cur = None          # current block: {'style':..,'runs':[]}
        self.b = 0; self.i = 0   # bold / italic depth
        self.skip = 0
        self.pages = 0

    def _anc(self, name):
        return any(name in c for (_, c) in self.stack)

    def _style(self, tag, cls):
        if tag == "p":
            if "big-quote" in cls: return "quote"
            if "say" in cls: return "say"
            if "bcap" in cls: return "bcap"
            if "cast-row" in cls: return "cast"
            if "tease" in cls: return "tease"
            if self._anc("case"): return "lesson"
            if self._anc("next"): return "nextdesc"
            if self._anc("theend"): return "theendp"
            if self._anc("disclaimer"): return "disclaimer"
            return "body"
        if tag == "h3":
            if self._anc("next"): return "nexttitle"
            if self._anc("theend"): return "theendtitle"
            return "body"
        return None

    def handle_starttag(self, tag, attrs):
        cls = set(dict(attrs).get("class", "").split())
        # scene break between page-sections
        if tag == "section" and "page" in cls:
            self.pages += 1
            if self.pages > 1 and self.cur is None:
                self.ops.append(("scene",))
        if tag == "img":
            src = dict(attrs).get("src", "")
            if src and self.skip == 0:
                self.ops.append(("image", os.path.basename(src)))
        self.stack.append((tag, cls))
        if self.skip:
            return
        if cls & SKIP:
            self.skip = len(self.stack); return
        if tag in ("strong", "b"): self.b += 1
        if tag in ("em", "i", "cite"): self.i += 1
        if self.cur is not None:
            return
        # open a block?
        style = self._style(tag, cls)
        if tag in ("span",) and ("ln" in cls or "blabel" in cls): style = "mono"
        elif tag == "div" and "bcell" in cls: style = "mono"
        elif tag == "span" and "nk" in cls: style = "nk"
        elif tag == "div" and "note" in cls: style = "note"
        elif tag == "span" and "tk" in cls: style = "tk"
        if style:
            self.cur = {"style": style, "runs": [], "depth": len(self.stack)}

    def handle_data(self, data):
        if self.skip or self.cur is None:
            return
        if data:
            self.cur["runs"].append((data, self.b > 0, self.i > 0))

    def handle_endtag(self, tag):
        depth = len(self.stack)
        if self.skip and depth == self.skip:
            self.skip = 0
        if self.cur and depth == self.cur["depth"]:
            text = "".join(r[0] for r in self.cur["runs"])
            if text.strip():
                self.ops.append(("para", self.cur["style"], self.cur["runs"]))
            self.cur = None
        if not self.skip:
            if tag in ("strong", "b") and self.b: self.b -= 1
            if tag in ("em", "i", "cite") and self.i: self.i -= 1
        if self.stack:
            self.stack.pop()

# ---------- docx helpers ----------
def add_runs(p, runs, italic=False, bold=False, mono=False, size=None, color=None):
    collapsed = re.sub(r"\s+", " ", "".join(r[0] for r in runs))  # for simple styles
    if not any(r[1] or r[2] for r in runs) and (italic or bold or mono):
        r = p.add_run(collapsed)
        r.italic = italic or None; r.bold = bold or None
        if mono: r.font.name = "Courier New"
        if size: r.font.size = Pt(size)
        if color: r.font.color.rgb = color
        return
    for text, b, i in runs:
        text = re.sub(r"\s+", " ", text)
        if not text: continue
        r = p.add_run(text)
        r.bold = (b or bold) or None
        r.italic = (i or italic) or None
        if mono: r.font.name = "Courier New"
        if size: r.font.size = Pt(size)
        if color: r.font.color.rgb = color

def field(paragraph, instr):
    """Insert a Word field (e.g. PAGE, or a TOC) into a paragraph."""
    run = paragraph.add_run()
    for t, txt in (("begin", None), ("instrText", instr), ("separate", None), ("text", "—"), ("end", None)):
        el = OxmlElement("w:" + ("fldChar" if t in ("begin", "separate", "end") else t))
        if t in ("begin", "separate", "end"):
            el.set(qn("w:fldCharType"), t)
        else:
            el.set(qn("xml:space"), "preserve"); el.text = txt
        run._r.append(el)

def render(doc, ops):
    for op in ops:
        if op[0] == "scene":
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("⁂"); r.font.size = Pt(12); r.font.color.rgb = GRAY
            continue
        if op[0] == "image":
            path = os.path.join(ROOT, "story", op[1])
            if os.path.exists(path):
                doc.add_picture(path, width=Inches(3.2))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        _, style, runs = op
        p = doc.add_paragraph()
        if style == "body":
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; add_runs(p, runs)
        elif style == "say":
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; add_runs(p, runs, italic=True)
        elif style == "quote":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_runs(p, runs, italic=True, size=13)
        elif style == "bcap":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_runs(p, runs, italic=True, size=9.5, color=GRAY)
        elif style in ("cast", "tease"):
            add_runs(p, runs); p.paragraph_format.left_indent = Inches(0.2)
        elif style == "lesson":
            r = p.add_run("✓  "); r.font.color.rgb = GREEN; add_runs(p, runs)
        elif style == "mono":
            add_runs(p, runs, mono=True, size=9)
            p.paragraph_format.left_indent = Inches(0.2); p.paragraph_format.space_after = Pt(0)
        elif style == "note":
            add_runs(p, runs, italic=True); p.paragraph_format.left_indent = Inches(0.3)
        elif style == "nk":
            add_runs(p, runs, size=9, color=GREEN)
        elif style == "nexttitle":
            add_runs(p, runs, bold=True, size=12)
        elif style == "nextdesc":
            add_runs(p, runs, italic=True, color=GRAY)
        elif style == "tk":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_runs(p, runs, size=9, color=GREEN)
        elif style == "theendtitle":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_runs(p, runs, bold=True, size=15)
        elif style == "theendp":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_runs(p, runs, italic=True)
        elif style == "disclaimer":
            add_runs(p, runs, italic=True, size=9.5, color=GRAY)
        else:
            add_runs(p, runs)

def main():
    doc = Document()
    # base style
    normal = doc.styles["Normal"]
    normal.font.name = "Georgia"; normal.font.size = Pt(11.5)
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.4
    # page size 5x8 + margins
    sec = doc.sections[0]
    sec.page_width = Inches(5); sec.page_height = Inches(8)
    sec.top_margin = sec.bottom_margin = Inches(0.7)
    sec.left_margin = sec.right_margin = Inches(0.65)
    # footer page number
    fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    field(fp, "PAGE")

    # --- front matter ---
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(120)
    r = t.add_run(META["title"]); r.bold = True; r.font.size = Pt(26)
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run(META["subtitle"]); r.italic = True; r.font.size = Pt(13); r.font.color.rgb = GRAY
    a = doc.add_paragraph(); a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    a.paragraph_format.space_before = Pt(48)
    a.add_run(META["author"])
    doc.add_page_break()

    for line in ("Copyright © 2026 %s. All rights reserved." % META["author"],
                 "First edition, 2026.", DISCLAIMER,
                 "No part of this book may be reproduced or used in any manner without the prior "
                 "written permission of the author, except for brief quotations in a review."):
        p = doc.add_paragraph(); r = p.add_run(line); r.font.size = Pt(9.5)
        if line is DISCLAIMER: r.italic = True
    doc.add_page_break()

    d = doc.add_paragraph(); d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.paragraph_format.space_before = Pt(180)
    r = d.add_run(DEDICATION); r.italic = True; r.font.size = Pt(12.5)
    doc.add_page_break()

    h = doc.add_heading("Contents", level=1)
    field(doc.add_paragraph(), 'TOC \\o "1-1" \\h \\z \\u')
    doc.add_page_break()

    # --- chapters ---
    files = sorted(glob.glob(os.path.join(ROOT, "_chapters", "*.html")))
    chapters = sorted((load_chapter(f) for f in files), key=lambda c: c["order"])
    for ch in chapters:
        head = doc.add_heading(level=1)
        head.paragraph_format.page_break_before = True
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if ch["num"]:
            run = head.add_run("Chapter %s\n" % ch["num"]); run.font.size = Pt(11); run.font.color.rgb = GREEN
        head.add_run(ch["title"])
        if ch["subtitle"]:
            sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = sp.add_run(ch["subtitle"]); r.italic = True; r.font.color.rgb = GRAY
        w = Walker(); w.feed(ch["body"]); w.close()
        render(doc, w.ops)

    # --- back matter ---
    doc.add_heading("The Testing Ideas in This Book", level=1).paragraph_format.page_break_before = True
    sp = doc.add_paragraph(); r = sp.add_run("A plain-English glossary, in alphabetical order."); r.italic = True; r.font.color.rgb = GRAY
    for term, desc in GLOSSARY:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8)
        rt = p.add_run(term + "  "); rt.bold = True
        p.add_run("— " + desc)

    doc.add_heading("Acknowledgements", level=1).paragraph_format.page_break_before = True
    for line in ACK:
        doc.add_paragraph(line)

    doc.add_heading("About the Author", level=1).paragraph_format.page_break_before = True
    doc.add_paragraph(BIO)
    lp = doc.add_paragraph(); lp.paragraph_format.space_before = Pt(12)
    r = lp.add_run("Connect on LinkedIn:  linkedin.com/in/hariprasadms"); r.font.color.rgb = GREEN

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print("wrote %s (%d KB)" % (OUT, os.path.getsize(OUT) // 1024))

if __name__ == "__main__":
    main()
